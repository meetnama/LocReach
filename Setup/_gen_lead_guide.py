"""Generate Setup/LocHere_Lead_Discovery_Guide.docx (LocReach setup & user guide)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = Path(__file__).resolve().parent / "LocHere_Lead_Discovery_Guide.docx"


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def h1(t: str) -> None:
        doc.add_heading(t, level=1)

    def h2(t: str) -> None:
        doc.add_heading(t, level=2)

    def h3(t: str) -> None:
        doc.add_heading(t, level=3)

    def para(t: str = ""):
        return doc.add_paragraph(t)

    def bullet(t: str) -> None:
        doc.add_paragraph(t, style="List Bullet")

    def numbered(t: str) -> None:
        doc.add_paragraph(t, style="List Number")

    title = doc.add_heading("LocHere Lead Discovery Guide", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        "Product UI name: LocReach  ·  Updated: 16 July 2026"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(
        "Authoritative companion for Sales_Tool / localization-leads. "
        "Matches the live 3-step Streamlit app. Launchers live in the Setup folder."
    )

    h1("1. What this tool does")
    para(
        "LocReach (LocHere Lead Discovery) finds B2B leads in localization / "
        "translation and related markets:"
    )
    numbered(
        "Find & Qualify Domains — search the web, scrape/score company sites, "
        "keep strong/possible matches."
    )
    numbered(
        "Find People — decision-maker contacts (project / vendor / translation managers)."
    )
    numbered(
        "Find Emails — confirmed work emails only (website crawl, verified patterns, search)."
    )
    para(
        "Flow is always: discover → qualify → export. "
        "The app does not import domain lists from files."
    )

    h1("2. What you need")
    bullet("Windows PC with Google Chrome")
    bullet("Python 3.11+ (install via Setup/1 - Get Python.bat — tick Add to PATH)")
    bullet("Docker Desktop (recommended) for SearXNG + OpenSERP free search backends")
    bullet("Folder layout: Sales_Tool\\Setup\\ and Sales_Tool\\localization-leads\\")

    h1("3. First-time setup (run in order)")
    numbered("1 - Get Python.bat")
    numbered("2 - Get Chrome.bat (skip if Chrome already installed)")
    numbered("3 - Get Docker.bat (optional but strongly recommended)")
    numbered(
        "4 - Install LocReach.bat — creates venv, installs packages, "
        "Desktop/Start Menu shortcuts"
    )
    numbered(
        "Start Docker Desktop, then 6 - Start LocReach.bat "
        "(auto-starts SearXNG/OpenSERP when possible)"
    )
    para(
        "Optional: 5 - Start SearXNG.bat and 7 - Start OpenSERP.bat "
        "if you prefer to start backends manually."
    )
    para(
        "To move to a new PC: see MOVE TO NEW PC.txt in this Setup folder. "
        "Copy the whole Sales_Tool folder (you may skip venv and .chrome_profile). "
        "Also copy leads.db and .env to keep data."
    )

    h1("4. Daily use")
    numbered("Open Docker Desktop (recommended).")
    numbered("Double-click 6 - Start LocReach.bat (or the Desktop shortcut).")
    numbered(
        "Browser opens http://localhost:8501 — use the top bar: "
        "Home · Step 1 · Step 2 · Step 3."
    )
    numbered(
        "When finished, close the LocReach browser tab so the watchdog can stop Streamlit."
    )
    para(
        "Tip: Minimizing Chrome during a long Step 1 run is fine. "
        "If Chrome Memory Saver sleeps localhost tabs, exclude http://localhost "
        "in chrome://settings/performance."
    )
    para(
        "If you see “Connection error — Is Streamlit still running?”, LocReach was "
        "stopped or crashed; that message is from the orphaned browser tab. "
        "Restart via bat 6 or close the tab."
    )

    h1("5. The 3-step pipeline")

    h2("Step 1 — Find & Qualify Domains")
    para("Page: Step 1 in the top nav. Settings live in the Search Settings expander.")
    bullet(
        "Pick Industry (e.g. Localization) and Country (e.g. Egypt), "
        "optional custom search term(s)."
    )
    bullet(
        "Set Target companies (total) — how many new qualified companies you want this run."
    )
    bullet(
        "Workers: 16 / 32 / 48 (default 32) — speeds scrape/qualify; "
        "search is often the bottleneck."
    )
    bullet("Press Find & Qualify Domains. Use Stop to halt early.")

    h3("What happens under the hood")
    bullet(
        "Search engines (free): SearXNG ∥ OpenSERP race → Google gap-fill → Bing → DuckDuckGo."
    )
    bullet("Cheap screen: drop blocked / duplicate / already-qualified domains.")
    bullet(
        "Scrape & score: keep strong and possible tiers "
        "(LinkedIn helps but is not required)."
    )
    bullet(
        "Search terms: primary list + expansion + auto-rotate bank of unused query templates."
    )
    bullet(
        "Stops when: target met, scrape budget full (~25× target), "
        "SERP yield flattens (diminishing returns), terms exhausted, or you press Stop."
    )

    h3("Reading results")
    bullet(
        "UI shows qualified count, checked, failed, unreachable, "
        "and a DB summary for your industry/country."
    )
    bullet(
        "Detailed engine/term diagnostics are NOT on the screen — open "
        r"localization-leads\logs\step1_search.log after a run."
    )
    bullet(
        "Export this run’s qualified domains with the Step 1 download button "
        "when the run finishes."
    )

    h3("Same market again (e.g. more Egypt LSPs)")
    para(
        "Re-running the same industry+country often returns many duplicates because "
        "already-qualified domains are skipped and search engines reshuffle the same "
        "top results. Auto-rotate helps, but for more coverage change keyword/country "
        "or raise the target so deeper templates run. For a clean benchmark wipe: stop "
        r"LocReach → delete localization-leads\leads.db → restart → re-run."
    )

    h2("Step 2 — Find People")
    bullet("Works on qualified domains from Step 1.")
    bullet("Classifies LSP vs client; finds people via search / LinkedIn / company site.")
    bullet(
        "Title filter keeps only: project manager, vendor manager, translation manager."
    )
    bullet(
        "Optional LinkedIn login via .env (LINKEDIN_EMAIL / LINKEDIN_PASSWORD) "
        "for /people scraping."
    )

    h2("Step 3 — Find Emails")
    bullet("L1: emails found on company websites (always kept if found).")
    bullet("L2: pattern + SMTP verify — only SMTP “good” results are saved.")
    bullet("L4: SearXNG people@domain search — confirmed when found.")
    bullet(
        "Unconfirmed guesses are not saved. SMTP needs dnspython (installed by bat 4); "
        "home networks often block port 25."
    )

    h1("6. Export")
    bullet("Each step can export the current run (Excel / HTML as shown on that page).")
    bullet("Full-DB Excel uses the unified domains + leads tables.")
    bullet(
        "Always prefer export out of LocReach — do not upload domain dumps into the tool."
    )

    h1("7. Important files")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Path"
    hdr[1].text = "Purpose"
    for a, b in [
        ("Setup/6 - Start LocReach.bat", "Daily launch (+ auto search backends)"),
        ("Setup/LocHere_Lead_Discovery_Guide.docx", "This user guide"),
        ("localization-leads/leads.db", "SQLite database (domains, people, leads)"),
        ("localization-leads/.env", "SEARXNG_URL, LinkedIn, Groq keys"),
        ("localization-leads/logs/step1_search.log", "Step 1 search diagnostics"),
        ("localization-leads/PROJECT_MEMORY.md", "Dev snapshot of how the app works"),
        ("searxng/settings.yml", "SearXNG config (must allow JSON API)"),
    ]:
        cells = table.add_row().cells
        cells[0].text = a
        cells[1].text = b

    h1("8. Troubleshooting")
    for title_t, fix in [
        (
            "Step 1 stops very early (e.g. ~13/100)",
            "Restart LocReach so check_budget ×25 is loaded; use bat 6.",
        ),
        (
            "Step 1 finishes in ~40s with almost no terms",
            "Bug is fixed in current code; update 1_Domains.py and restart. "
            "Log should list many TERM lines.",
        ),
        (
            "OpenSERP “results is not defined”",
            "Fixed in sources/utils.py; restart the app.",
        ),
        (
            "CAPTCHA / rate limits",
            "Keep Docker SearXNG + OpenSERP up; wait out Google cooldown; try later.",
        ),
        (
            "Excel export fails",
            "Update code (domains table export); re-run Install LocReach if needed.",
        ),
        (
            "Step 3 finds no emails",
            "Install dnspython via bat 4; expect L1/L4 if SMTP port 25 is blocked.",
        ),
        (
            "App dies mid-scan",
            "Don’t close the LocReach tab mid-run; 45s watchdog; Memory Saver exclusions.",
        ),
        (
            "Chrome / LinkedIn weird",
            r"Delete localization-leads\.chrome_profile\ and sign in again once.",
        ),
        (
            "Clean Step 1 re-test",
            "Stop app → delete leads.db → start bat 6 → run again.",
        ),
    ]:
        h3(title_t)
        para(fix)

    h1("9. Tips for volume")
    bullet("Healthy Docker search backends matter more than raising workers alone.")
    bullet("Workers speed scrape/qualify; SERP feed is usually the bottleneck.")
    bullet(
        "~1000 qualified/hour is an enterprise goal — may need paid SERP later "
        "if free engines dry out."
    )
    bullet(
        "After a successful market run, export; then change angle "
        "(city keyword, auto-rotate bank, neighbor country) rather than expecting "
        "+500 new from the identical SERP."
    )

    h1("10. Version note")
    para(
        "This guide matches the product as of 16 July 2026: unified 3-step LocReach app, "
        "Step 1 auto-rotate + diminishing-returns stop, step1_search.log for diagnostics, "
        "no domain-list import. Older docs describing a 4-step pipeline "
        "(Domain Discovery → Site Scanner → People → Email) are obsolete. "
        "Windows launchers live in Sales_Tool\\Setup\\."
    )

    doc.save(OUT)
    print(f"Wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
